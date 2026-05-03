from configparser import ConfigParser
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def val(section, key):
    return section.get(key, "-") or "-"


def set_cell_text(cell, text, bold=False, size=10, rotate=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if rotate:
        tc_pr = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        text_direction = OxmlElement("w:textDirection")
        text_direction.set(qn("w:val"), "btLr")
        tc_pr.append(text_direction)


def qn(tag):
    from docx.oxml.ns import qn as docx_qn
    return docx_qn(tag)


def build_doc(config_path: Path, output_path: Path):
    parser = ConfigParser()
    parser.read(config_path, encoding="utf-8")

    meta = parser["meta"]
    values = parser["values"]

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Pt(24)
    section.right_margin = Pt(24)
    section.top_margin = Pt(24)
    section.bottom_margin = Pt(24)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("1. O'QUV ISHLARI")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    metrics = [
        ("Ma'ruza", "maruza", "maruza_jami", "maruza_amalda"),
        ("Amaliy mashg'ulot", "amaliy_mashgulot", "amaliy_mashgulot_jami", "amaliy_mashgulot_amalda"),
        ("Laboratoriya ishi", "laboratoriya_ishi", "laboratoriya_ishi_jami", "laboratoriya_ishi_amalda"),
        ("Maslahat", "maslahat", "maslahat_jami", "maslahat_amalda"),
        ("Nazorat", "nazorat", "nazorat_jami", "nazorat_amalda"),
        ("Taqrizlar", "taqrizlar", "taqrizlar_jami", "taqrizlar_amalda"),
        ("Kurs ishi", "kurs_ishi", "kurs_ishi_jami", "kurs_ishi_amalda"),
        ("Bitiruv ishi", "bitiruv_ishi", "bitiruv_ishi_jami", "bitiruv_ishi_amalda"),
        ("DAK(BMI) rahbarligi", "dak_bmi_rahbarligi", "dak_bmi_rahbarligi_jami", "dak_bmi_rahbarligi_amalda"),
        ("Amaliyot", "amaliyot", "amaliyot_jami", "amaliyot_amalda"),
        ("ITI", "iti", "iti_jami", "iti_amalda"),
        ("BMI ga taqriz", "bmi_ga_taqriz", "bmi_ga_taqriz_jami", "bmi_ga_taqriz_amalda"),
        ("Qayta topshirish", "qayta_topshirish", "qayta_topshirish_jami", "qayta_topshirish_amalda"),
        ("To'plangan reyting bali", "rating_ball", "rating_ball_jami", "rating_ball_amalda"),
    ]

    rows = 6
    cols = 2 + len(metrics)
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(0, cols - 1))

    set_cell_text(table.cell(0, 0), "№", bold=True, size=11)
    set_cell_text(table.cell(0, 1), "O'quv ishlari nomi,\nfakultet, gurux,\ntalabalar soni", bold=True, size=10)
    set_cell_text(table.cell(0, 2), "O'quv ishlari turi (soat)", bold=True, size=11)

    for index, (label, *_keys) in enumerate(metrics, start=2):
        set_cell_text(table.cell(1, index), label, bold=True, size=9, rotate=True)

    table.cell(2, 0).merge(table.cell(2, cols - 1))
    set_cell_text(table.cell(2, 0), f"{meta.get('semester_text', 'Semestr')} semestr", bold=True, size=11)

    description = (
        f"{meta.get('work_name', '-')}\n"
        f"({meta.get('faculty', '-')}, {meta.get('group_name', '-')}, "
        f"{meta.get('student_count', '-')} talaba)"
    )
    set_cell_text(table.cell(3, 0), "1.", size=11)
    set_cell_text(table.cell(3, 1), description, bold=True, size=10)

    for index, (_label, main_key, _jami_key, _amalda_key) in enumerate(metrics, start=2):
        set_cell_text(table.cell(3, index), val(values, main_key), size=11)

    table.cell(4, 0).merge(table.cell(5, 0))
    set_cell_text(table.cell(4, 0), "Jami", bold=True, size=11)
    set_cell_text(table.cell(4, 1), "Reja", bold=True, size=11)
    set_cell_text(table.cell(5, 1), "Amalda", bold=True, size=11)

    for index, (_label, _main_key, jami_key, amalda_key) in enumerate(metrics, start=2):
        set_cell_text(table.cell(4, index), val(values, jami_key), bold=True, size=11)
        set_cell_text(table.cell(5, index), val(values, amalda_key), bold=True, size=11)

    document.save(output_path)


if __name__ == "__main__":
    build_doc(Path(sys.argv[1]), Path(sys.argv[2]))
